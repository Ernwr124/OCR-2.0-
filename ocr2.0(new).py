import os
import re
import json
import requests
import traceback
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
import uuid
import threading
import time
from collections import Counter
from io import BytesIO

from pdf2image import convert_from_path
from PIL import Image, ImageEnhance, ImageFilter
import numpy as np
import cv2
import pytesseract
from flask import Flask, request, jsonify, send_file, render_template_string
from difflib import SequenceMatcher
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Imports from updates ---
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import re
import json
import cv2
import numpy as np
import pytesseract
from pdf2image import convert_from_path
import re
from datetime import datetime
import ollama
from PIL import Image
import io
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import threading
import time
import uuid
# --- End imports from updates ---


# --- Настройки ---
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
HISTORY_FOLDER = 'history'
# --- Updates to UPLOAD_FOLDER and RESULTS_FOLDER ---
RESULTS_FOLDER = 'results' # Defined in updates
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(HISTORY_FOLDER, exist_ok=True)
os.makedirs(RESULTS_FOLDER, exist_ok=True) # Created from updates


DPI = 400  # Увеличен для лучшего качества
USE_PREPROCESSING = True # This seems to be unused in the merged code, but kept from original
OLLAMA_MODEL = "qwen2.5:3b-instruct-q4_K_M" # Updated from 'phi3.5:3.8b-mini-instruct-q4_K_M' to 'phi3.5'
OLLAMA_API_URL = "http://localhost:11434/api/chat" # This is kept from original, but OLLAMA_MODEL is used in `extract_contract_data_with_ai` directly

REQUIRED_FIELDS = [
    "document_type", "contract_number", "sign_date", "expiry_date",
    "seller", "buyer", "amount", "currency",
    "validation_status", "extraction_accuracy"
]

PROCESSING_MESSAGES = [
    "Инициализация ИИ-движка...",
    "Конвертация PDF в высококачественные изображения...",
    "Анализ структуры и макета документа...",
    "Многоуровневое распознавание текста (OCR)...",
    "Применение контекстного анализа...",
    "Исправление ошибок распознавания...",
    "Извлечение ключевых полей с валидацией...",
    "Кросс-проверка данных между страницами...",
    "Формирование структурированного JSON...",
    "Расчет метрик качества и точности...",
    "Финализация результатов..."
]

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['HISTORY_FOLDER'] = HISTORY_FOLDER
# --- Updates to app config ---
app.config['RESULTS_FOLDER'] = RESULTS_FOLDER # Added from updates
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size # Added from updates

# --- Global storage for processing queue and history ---
processing_queue = {}
analysis_history = []
# --- End global storage ---


# --- УЛУЧШЕННАЯ АРХИТЕКТУРА OCR ---

class AdvancedOCREngine:
    """
    Продвинутый OCR-движок с многоуровневой обработкой и контекстным пониманием
    """
    
    @staticmethod
    def advanced_preprocess_image(image: Image.Image, page_num: int) -> List[Image.Image]:
        """
        Создает несколько вариантов предобработки для повышения точности
        """
        variants = []
        
        # Конвертация в grayscale
        if image.mode != 'L':
            image = image.convert('L')
        
        # Вариант 1: Адаптивная бинаризация
        img_array = np.array(image)
        binary = cv2.adaptiveThreshold(
            img_array, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
            cv2.THRESH_BINARY, 21, 11
        )
        variants.append(Image.fromarray(binary))
        
        # Вариант 2: Улучшение контраста + резкость
        enhancer = ImageEnhance.Contrast(image)
        enhanced = enhancer.enhance(2.8)
        enhancer = ImageEnhance.Sharpness(enhanced)
        enhanced = enhancer.enhance(3.5)
        variants.append(enhanced)
        
        # Вариант 3: Удаление шума + морфологические операции
        denoised = cv2.fastNlMeansDenoising(img_array, None, 10, 7, 21)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        morph = cv2.morphologyEx(denoised, cv2.MORPH_CLOSE, kernel)
        variants.append(Image.fromarray(morph))
        
        # Вариант 4: Коррекция наклона
        coords = np.column_stack(np.where(img_array < 200))
        if len(coords) > 0:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            if abs(angle) > 0.5:
                (h, w) = img_array.shape[:2]
                center = (w // 2, h // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated = cv2.warpAffine(
                    img_array, M, (w, h), 
                    flags=cv2.INTER_CUBIC, 
                    borderMode=cv2.BORDER_REPLICATE
                )
                variants.append(Image.fromarray(rotated))
        
        return variants
    
    @staticmethod
    def extract_text_multi_variant(image: Image.Image, page_num: int) -> str:
        """
        Извлекает текст используя несколько вариантов предобработки и выбирает лучший
        """
        variants = AdvancedOCREngine.advanced_preprocess_image(image, page_num)
        
        results = []
        for idx, variant in enumerate(variants):
            try:
                # Используем разные конфигурации Tesseract
                configs = [
                    '--psm 1 --oem 3',  # Автоматическая сегментация
                    '--psm 3 --oem 3',  # Полностью автоматическая
                    '--psm 6 --oem 3',  # Единый блок текста
                ]
                
                for config in configs:
                    text = pytesseract.image_to_string(
                        variant, 
                        lang='rus+eng',
                        config=config
                    )
                    if text.strip():
                        results.append(text.strip())
            except Exception as e:
                print(f"  ⚠️ Ошибка в варианте {idx}: {e}")
                continue
        
        # Выбираем самый длинный и информативный результат
        if results:
            best_result = max(results, key=lambda x: len(x))
            return best_result
        return ""
    
    @staticmethod
    def extract_with_context(text: str, field_type: str) -> Optional[str]:
        """
        Извлекает поля с учетом контекста и окружающих слов
        """
        patterns = {
            'contract_number': [
                r'(?:Контракт|CONTRACT|Договор|ДОГОВОР)\s*[№#No\.]*\s*([-\w\/\.\-]+)',
                r'№\s*([-\w\/\.\-]+)\s*(?:от|from)',
                r'ДОГОВОР\s+ПОСТАВКИ\s*№\s*([-\w\/\.\-]+)',
            ],
            'sign_date': [
                r'(?:от|from|dated)\s*[«"]?(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})[»"]?',
                r'(\d{1,2}\.\d{1,2}\.\d{4})\s*(?:года|г\.)?',
                r'[«"](\d{1,2})[»"]\s*\w+\s*(\d{4})',
            ],
            'expiry_date': [
                r'(?:до|until|по)\s*[«"]?(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})[»"]?',
                r'(?:срок действия|validity).*?(\d{1,2}\.\d{1,2}\.\d{4})',
                r'(?:действует до|valid until)\s*[«"]?(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})[»"]?',
            ],
            'seller': [
                r'(?:Продавец|Seller|ПРОДАВЕЦ|SELLER)[:\s,]+([^,\n]+?)(?:,\s*в\s*лице|,\s*именуемое|represented)',
                r'(?:ТОО|LLP|ООО|LLC)\s+[«"]?([^»"\n,]+)[»"]?',
            ],
            'buyer': [
                r'(?:Покупатель|Buyer|ПОКУПАТЕЛЬ|BUYER)[:\s,]+([^,\n]+?)(?:,\s*в\s*лице|,\s*именуемое|represented)',
                r'(?:UAB|ТОО|LLP|ООО|LLC)\s+[«"]?([^»"\n,]+)[»"]?',
            ],
            'amount': [
                r'(?:общая\s+сумма|total\s+amount|сумма\s+контракта|contract\s+amount)[:\s]+(\d[\d\s,\.]+)',
                r'(?:на\s+сумму|for\s+the\s+amount)[:\s]+(\d[\d\s,\.]+)',
                r'(?:стоимость|cost|цена|price)[:\s]+(\d[\d\s,\.]+)',
                r'(\d[\d\s,\.]+)\s*(?:USD|EUR|RUB|KZT)',
                r'(?:сумма|amount)[:\s]+(\d[\d\s,\.]+)',
                r'(?:всего|total)[:\s]+(\d[\d\s,\.]+)',
                r'(\d{1,3}(?:[\s,]\d{3})*(?:\.\d{2})?)\s*(?:долларов|евро|рублей|тенге)',
                r'(?:составляет|amounts\s+to)[:\s]+(\d[\d\s,\.]+)',
            ],
            'currency': [
                r'(USD|EUR|RUB|KZT|долларов|доллар|евро|рублей|рубль|тенге)',
                r'в\s+(долларах|евро|рублях|тенге)',
            ],
        }
        
        if field_type not in patterns:
            return None
        
        for pattern in patterns[field_type]:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                value = match.group(1).strip()
                if value and len(value) > 1:
                    return value
        
        return None
    
    @staticmethod
    def validate_and_clean_field(value: Optional[str], field_type: str) -> Optional[str]:
        """
        Валидирует и очищает извлеченное поле
        """
        if not value:
            return None
        
        # Очистка от лишних символов
        value = re.sub(r'[\n\t\r]+', ' ', value).strip()
        value = re.sub(r'\s+', ' ', value)
        
        if field_type == 'contract_number':
            # Удаляем лишние символы из номера договора
            value = re.sub(r'[^\w\/\-\.]', '', value)
            return value if len(value) > 2 else None
        
        elif field_type in ['sign_date', 'expiry_date']:
            # Нормализация формата даты
            date_match = re.search(r'(\d{1,2})[\.\/\-](\d{1,2})[\.\/\-](\d{2,4})', value)
            if date_match:
                day, month, year = date_match.groups()
                if len(year) == 2:
                    year = '20' + year
                try:
                    # Проверка на корректность даты
                    datetime.strptime(f"{day}.{month}.{year}", "%d.%m.%Y")
                    return f"{day.zfill(2)}.{month.zfill(2)}.{year}"
                except ValueError:
                    return None
            return None
        
        elif field_type in ['seller', 'buyer']:
            # Очистка названий компаний
            value = re.sub(r'[«»""\'\[\]]', '', value)
            value = value.split(',')[0].strip()  # Берем только название до запятой
            return value if len(value) > 3 else None
        
        elif field_type == 'amount':
            value = re.sub(r'[^\d,\.\s]', '', value)
            value = value.replace(' ', '').replace(',', '.')
            # Удаляем лишние точки, оставляем только последнюю
            parts = value.split('.')
            if len(parts) > 2:
                value = ''.join(parts[:-1]) + '.' + parts[-1]
            try:
                float_val = float(value)
                if float_val > 0:
                    return value
            except ValueError:
                pass
            return None
        
        elif field_type == 'currency':
            # Нормализация валюты
            currency_map = {
                'долларов': 'USD', 'доллар': 'USD',
                'евро': 'EUR',
                'рублей': 'RUB', 'рубль': 'RUB',
                'тенге': 'KZT',
            }
            value_lower = value.lower()
            for key, val in currency_map.items():
                if key in value_lower:
                    return val
            if value.upper() in ['USD', 'EUR', 'RUB', 'KZT']:
                return value.upper()
            return None
        
        return value

class IntelligentDataExtractor:
    """
    Интеллектуальный экстрактор данных с использованием LLM
    """
    
    @staticmethod
    def call_llm_with_retry(extracted_text: str, page_num: int, max_retries: int = 2) -> Dict[str, Any]:
        """
        Вызывает LLM с повторными попытками
        """
        prompt = f"""Ты — эксперт-аналитик банковских документов. Извлеки ТОЧНЫЕ данные из договора.

КРИТИЧЕСКИ ВАЖНО - ищи эти поля:
1. contract_number - номер после "Контракт №", "Договор №", "CONTRACT №"
2. sign_date - дата после "от", "dated" (формат ДД.ММ.ГГГГ)
3. expiry_date - дата после "действует до", "до", "until"
4. seller - название после "Продавец:", "Seller:"
5. buyer - название после "Покупатель:", "Buyer:"
6. amount - ОБЯЗАТЕЛЬНО найди сумму! Ищи "общая сумма", "сумма контракта", "на сумму", "стоимость", "всего". Только цифры!
7. currency - USD, EUR, RUB, KZT
8. summary - Краткое описание договора на 70-80 слов: предмет договора, стороны, основные условия

Верни ТОЛЬКО JSON:
{{"contract_number": "...", "sign_date": "ДД.ММ.ГГГГ", "expiry_date": "ДД.ММ.ГГГГ", 
  "seller": "...", "buyer": "...", "amount": "123456.78", "currency": "USD",
  "summary": "Краткое описание договора..."}}

Если поле не найдено - ставь null.

ТЕКСТ:
{extracted_text[:3000]}
"""
        
        for attempt in range(max_retries):
            try:
                payload = {
                    "model": OLLAMA_MODEL,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                    "options": {"temperature": 0.05, "top_p": 0.9}
                }
                
                response = requests.post(OLLAMA_API_URL, json=payload, timeout=120)
                response.raise_for_status()
                result = response.json()
                raw_content = result["message"]["content"].strip()
                
                json_match = re.search(r'\{.*\}', raw_content, re.DOTALL)
                if json_match:
                    json_string = json_match.group(0)
                    data = json.loads(json_string)
                    return data
                    
            except Exception as e:
                print(f"  ⚠️ Попытка {attempt + 1} не удалась: {e}")
                if attempt < max_retries - 1:
                    time.sleep(1)
                continue
        
        return {}
    
    @staticmethod
    def merge_results(llm_results: List[Dict], regex_results: Dict, all_text: str) -> Dict[str, Any]:
        """
        Объединяет результаты LLM и regex, выбирая наиболее достоверные данные
        """
        final_data = {
            "contract_number": None,
            "sign_date": None,
            "expiry_date": None,
            "seller": None,
            "buyer": None,
            "amount": None,
            "currency": None,
            "summary": None,  # Добавлено поле summary
        }
        
        # Собираем все варианты для каждого поля
        field_candidates = {key: [] for key in final_data.keys()}
        
        # Добавляем результаты LLM
        for result in llm_results:
            if isinstance(result, dict):
                for key in final_data.keys():
                    if key in result and result[key]:
                        if key == 'summary':
                            # Для summary берем первое непустое значение
                            field_candidates[key].append(str(result[key]))
                        else:
                            cleaned = AdvancedOCREngine.validate_and_clean_field(
                                str(result[key]), key
                            )
                            if cleaned:
                                field_candidates[key].append(cleaned)
        
        # Добавляем результаты regex (кроме summary)
        for key, value in regex_results.items():
            if value and key != 'summary':
                field_candidates[key].append(value)
        
        # Выбираем наиболее частый вариант (консенсус)
        for key, candidates in field_candidates.items():
            if candidates:
                if key == 'summary':
                    # Для summary берем самое длинное описание
                    final_data[key] = max(candidates, key=len)
                else:
                    # Подсчитываем частоту каждого варианта
                    counter = Counter(candidates)
                    most_common = counter.most_common(1)[0][0]
                    final_data[key] = most_common
        
        return final_data

def extract_structured_data_v2(
    pages: List[Image.Image], 
    filename: str,
    status_callback=None
) -> Dict[str, Any]:
    """
    Улучшенная функция извлечения данных с многоуровневым подходом
    """
    final_data = {
        "document_type": "contract",
        "file_name": filename,
        "contract_number": None,
        "sign_date": None,
        "expiry_date": None,
        "seller": None,
        "buyer": None,
        "amount": None,
        "currency": None,
        "validation_status": "partial",
        "extraction_accuracy": 0.0,
        "confidence_scores": {},
        "metrics": {
            "CER": 0.0,
            "WER": 0.0,
            "Levenshtein": 0.0,
            "field_level_accuracy": 0.0,
            "exact_match": 0.0,
            "json_validity": 0.0,
            "schema_consistency": 0.0,
        },
        "summary": None # Added for summary
    }
    
    llm_results = []
    all_pages_text = ""
    
    # Шаг 1: OCR всех страниц
    for i, page in enumerate(pages):
        if status_callback:
            status_callback(f"Обработка страницы {i+1}/{len(pages)}...", 20 + (i / len(pages)) * 30)
        
        print(f"  📄 Страница {i+1}: Многовариантное распознавание...")
        page_text = AdvancedOCREngine.extract_text_multi_variant(page, i+1)
        all_pages_text += f"\n--- Страница {i+1} ---\n{page_text}"
        
        # Вызываем LLM для каждой страницы
        if status_callback:
            status_callback(f"ИИ-анализ страницы {i+1}/{len(pages)}...", 50 + (i / len(pages)) * 20)
        
        llm_result = IntelligentDataExtractor.call_llm_with_retry(page_text, i+1)
        llm_results.append(llm_result)
    
    # Шаг 2: Извлечение с помощью regex
    if status_callback:
        status_callback("Контекстный анализ всего документа...", 75)
    
    regex_results = {}
    for field in ['contract_number', 'sign_date', 'expiry_date', 'seller', 'buyer', 'amount', 'currency']:
        extracted = AdvancedOCREngine.extract_with_context(all_pages_text, field)
        cleaned = AdvancedOCREngine.validate_and_clean_field(extracted, field)
        regex_results[field] = cleaned
    
    # Шаг 3: Объединение результатов
    if status_callback:
        status_callback("Кросс-валидация данных...", 85)
    
    merged_data = IntelligentDataExtractor.merge_results(llm_results, regex_results, all_pages_text)
    
    # Обновляем финальные данные
    for key in ['contract_number', 'sign_date', 'expiry_date', 'seller', 'buyer', 'amount', 'currency']:
        final_data[key] = merged_data[key]
    
    # Добавляем краткое описание, если оно есть в LLM результатах
    # Сначала ищем в объединенном результате, затем в LLM результатах
    final_data['summary'] = merged_data.get('summary')
    if not final_data['summary']:
        for result in llm_results:
            if isinstance(result, dict) and 'summary' in result and result['summary']:
                final_data['summary'] = result['summary']
                break # Берем первое найденное описание
    
    # Шаг 4: Расчет метрик
    if status_callback:
        status_callback("Расчет метрик качества...", 95)
    
    total_fields = 7
    found_fields = sum(1 for key in ['contract_number', 'sign_date', 'expiry_date', 'seller', 'buyer', 'amount', 'currency'] 
                      if final_data.get(key) is not None)
    
    final_data["extraction_accuracy"] = round(found_fields / total_fields, 2)
    final_data["validation_status"] = "valid" if found_fields >= 6 else "partial"
    final_data["metrics"]["field_level_accuracy"] = final_data["extraction_accuracy"]
    final_data["metrics"]["exact_match"] = 1.0 if found_fields == total_fields else 0.0
    
    valid_llm = sum(1 for r in llm_results if isinstance(r, dict) and len(r) > 0)
    final_data["metrics"]["json_validity"] = valid_llm / len(llm_results) if llm_results else 0.0
    final_data["metrics"]["schema_consistency"] = final_data["metrics"]["json_validity"]
    
    # Оценка уверенности для каждого поля
    for key in ['contract_number', 'sign_date', 'expiry_date', 'seller', 'buyer', 'amount', 'currency']:
        if final_data[key]:
            # Подсчитываем, сколько источников согласны с этим значением
            agreement_count = 0
            # Проверяем LLM результаты
            for result in llm_results:
                if isinstance(result, dict) and result.get(key) == final_data[key]:
                    agreement_count += 1
            # Проверяем Regex результаты
            if regex_results.get(key) == final_data[key]:
                agreement_count += 1
            
            # Уверенность рассчитывается на основе количества источников, согласных с финальным значением
            # Источники: LLM результаты (len(llm_results)) + Regex результат (1)
            max_possible_agreement = len(llm_results) + 1
            confidence = min(agreement_count / max_possible_agreement, 1.0) if max_possible_agreement > 0 else 0.0
            final_data["confidence_scores"][key] = round(confidence, 2)
    
    return final_data

def save_to_history(data: Dict[str, Any], analysis_id: str):
    """Сохраняет результат анализа в историю"""
    history_file = os.path.join(HISTORY_FOLDER, f"{analysis_id}.json")
    data['analysis_id'] = analysis_id
    data['timestamp'] = datetime.now().isoformat()
    
    # Убедимся, что 'extracted_data' не хранится отдельно, если data уже содержит всё
    if 'extracted_data' in data:
        del data['extracted_data'] # Avoid redundancy if data is already the final output

    with open(history_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_history() -> List[Dict[str, Any]]:
    """Загружает историю анализов"""
    history = []
    if not os.path.exists(HISTORY_FOLDER):
        os.makedirs(HISTORY_FOLDER)
        
    for filename in os.listdir(HISTORY_FOLDER):
        if filename.endswith('.json'):
            try:
                with open(os.path.join(HISTORY_FOLDER, filename), 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    history.append(data)
            except Exception as e:
                print(f"Ошибка загрузки истории {filename}: {e}")
    
    # Сортируем по времени (новые первыми)
    history.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
    return history

def export_to_excel(data):
    """
    Экспортирует данные в Excel файл с профессиональным форматированием
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Анализ договора"
    
    # Заголовок
    ws['A1'] = 'АНАЛИЗ БАНКОВСКОГО ДОГОВОРА'
    ws['A1'].font = Font(size=16, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='1E3A8A', end_color='1E3A8A', fill_type='solid')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:B1')
    ws.row_dimensions[1].height = 30
    
    # Общая информация
    row = 3
    ws[f'A{row}'] = 'ОБЩАЯ ИНФОРМАЦИЯ'
    ws[f'A{row}'].font = Font(size=12, bold=True, color='FFFFFF')
    ws[f'A{row}'].fill = PatternFill(start_color='3B82F6', end_color='3B82F6', fill_type='solid')
    ws.merge_cells(f'A{row}:B{row}')
    
    row += 1
    # The 'data' parameter directly contains the structured data from extract_structured_data_v2
    extracted = data 
    
    info_data = [
        ('Тип документа', extracted.get('document_type', 'contract')),
        ('Имя файла', extracted.get('file_name', '')),
        ('Статус валидации', extracted.get('validation_status', 'valid')),
    ]
    
    for label, value in info_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
    
    # Краткое описание
    row += 1
    ws[f'A{row}'] = 'КРАТКОЕ ОПИСАНИЕ ДОГОВОРА'
    ws[f'A{row}'].font = Font(size=12, bold=True, color='FFFFFF')
    ws[f'A{row}'].fill = PatternFill(start_color='10B981', end_color='10B981', fill_type='solid')
    ws.merge_cells(f'A{row}:B{row}')
    
    row += 1
    summary = extracted.get('summary', 'Описание недоступно')
    ws[f'A{row}'] = summary
    ws[f'A{row}'].alignment = Alignment(wrap_text=True, vertical='top')
    ws.merge_cells(f'A{row}:B{row}')
    ws.row_dimensions[row].height = 80
    
    # Данные договора
    row += 2
    ws[f'A{row}'] = 'ДАННЫЕ ДОГОВОРА'
    ws[f'A{row}'].font = Font(size=12, bold=True, color='FFFFFF')
    ws[f'A{row}'].fill = PatternFill(start_color='3B82F6', end_color='3B82F6', fill_type='solid')
    ws.merge_cells(f'A{row}:B{row}')
    
    row += 1
    contract_data = [
        ('Номер договора', extracted.get('contract_number', '—')),
        ('Дата подписания', extracted.get('sign_date', '—')),
        ('Дата окончания', extracted.get('expiry_date', '—')),
        ('Продавец', extracted.get('seller', '—')),
        ('Покупатель', extracted.get('buyer', '—')),
        ('Сумма', extracted.get('amount', '—')),
        ('Валюта', extracted.get('currency', '—')),
    ]
    
    for label, value in contract_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
    
    # Метрики качества
    row += 1
    ws[f'A{row}'] = 'МЕТРИКИ КАЧЕСТВА'
    ws[f'A{row}'].font = Font(size=12, bold=True, color='FFFFFF')
    ws[f'A{row}'].fill = PatternFill(start_color='3B82F6', end_color='3B82F6', fill_type='solid')
    ws.merge_cells(f'A{row}:B{row}')
    
    row += 1
    metrics = extracted.get('metrics', {})
    metrics_data = [
        ('Точность извлечения', f"{metrics.get('field_level_accuracy', 0)*100:.1f}%"),
        ('Полное совпадение', f"{metrics.get('exact_match', 0)*100:.1f}%"),
        ('Валидность JSON', f"{metrics.get('json_validity', 0)*100:.1f}%"),
        ('Согласованность схемы', f"{metrics.get('schema_consistency', 0)*100:.1f}%"),
    ]
    
    for label, value in metrics_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
    
    # Форматирование колонок
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 50
    
    # Границы для всех ячеек
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for r in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=2):
        for cell in r:
            cell.border = thin_border
    
    # Сохранение
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return output

def export_to_docx(data):
    """
    Экспортирует данные в DOCX файл с профессиональным форматированием
    """
    doc = Document()
    
    # Заголовок документа
    heading = doc.add_heading('АНАЛИЗ БАНКОВСКОГО ДОГОВОРА', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.runs[0]
    heading_run.font.color.rgb = RGBColor(30, 58, 138)
    heading_run.font.size = Pt(24)
    
    doc.add_paragraph()
    
    # Общая информация
    section_heading = doc.add_heading('ОБЩАЯ ИНФОРМАЦИЯ', level=1)
    section_heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)
    
    extracted = data
    
    info_data = [
        ('Тип документа', extracted.get('document_type', 'contract')),
        ('Имя файла', extracted.get('file_name', '')),
        ('Статус валидации', extracted.get('validation_status', 'valid')),
    ]
    
    for label, value in info_data:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))
    
    doc.add_paragraph()
    
    # Краткое описание
    section_heading = doc.add_heading('КРАТКОЕ ОПИСАНИЕ ДОГОВОРА', level=1)
    section_heading.runs[0].font.color.rgb = RGBColor(16, 185, 129)
    
    summary = extracted.get('summary', 'Описание недоступно')
    summary_p = doc.add_paragraph(summary)
    summary_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Данные договора
    section_heading = doc.add_heading('ДАННЫЕ ДОГОВОРА', level=1)
    section_heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)
    
    contract_data = [
        ('Номер договора', extracted.get('contract_number', '—')),
        ('Дата подписания', extracted.get('sign_date', '—')),
        ('Дата окончания', extracted.get('expiry_date', '—')),
        ('Продавец', extracted.get('seller', '—')),
        ('Покупатель', extracted.get('buyer', '—')),
        ('Сумма', extracted.get('amount', '—')),
        ('Валюта', extracted.get('currency', '—')),
    ]
    
    for label, value in contract_data:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))
    
    doc.add_paragraph()
    
    # Метрики качества
    section_heading = doc.add_heading('МЕТРИКИ КАЧЕСТВА', level=1)
    section_heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)
    
    metrics = extracted.get('metrics', {})
    metrics_data = [
        ('Точность извлечения', f"{metrics.get('field_level_accuracy', 0)*100:.1f}%"),
        ('Полное совпадение', f"{metrics.get('exact_match', 0)*100:.1f}%"),
        ('Валидность JSON', f"{metrics.get('json_validity', 0)*100:.1f}%"),
        ('Согласованность схемы', f"{metrics.get('schema_consistency', 0)*100:.1f}%"),
    ]
    
    for label, value in metrics_data:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))
    
    doc.add_paragraph()
    
    # Оценки уверенности ИИ
    section_heading = doc.add_heading('ОЦЕНКИ УВЕРЕННОСТИ ИИ', level=1)
    section_heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)
    
    confidence_scores = extracted.get('confidence_scores', {})
    confidence_labels = {
        'contract_number': 'Номер договора',
        'sign_date': 'Дата подписания',
        'expiry_date': 'Дата окончания',
        'seller': 'Продавец',
        'buyer': 'Покупатель',
        'amount': 'Сумма',
        'currency': 'Валюта'
    }
    
    for key, value in confidence_scores.items():
        label = confidence_labels.get(key, key)
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(f'{value * 100:.0f}%')
    
    # Сохранение в BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

def create_excel_file(data: Dict[str, Any]) -> BytesIO:
    """Создает Excel файл с результатами анализа"""
    # В этой функции теперь вызывается новая, улучшенная функция export_to_excel
    return export_to_excel(data)

# --- Flask Routes ---

@app.route('/')
def index():
    # Original: return render_template_string(HTML_TEMPLATE)
    # Update implies reading from a file. For consistency with the provided HTML,
    # we'll use render_template_string with the inline HTML from the updates.
    return render_template_string(HTML_TEMPLATE) 


@app.route('/upload', methods=['POST'])
def upload_file():
    """Обработка загрузки нескольких PDF файлов"""
    if 'files' not in request.files:
        return jsonify({'error': 'Файлы не найдены'}), 400
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Файлы не выбраны'}), 400
    
    tasks = []
    
    for file in files:
        if file and file.filename.endswith('.pdf'):
            unique_filename = str(uuid.uuid4()) + '.pdf'
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            file.save(file_path)
            
            output_json_filename = str(uuid.uuid4()) + '.json'
            output_json_path = os.path.join(app.config['OUTPUT_FOLDER'], output_json_filename)
            
            analysis_id = str(uuid.uuid4())
            
            thread = threading.Thread(
                target=process_pdf_v2, 
                args=(file_path, output_json_path, unique_filename, output_json_filename, analysis_id, file.filename)
            )
            thread.start()
            
            tasks.append({
                'task_id': unique_filename,
                'output_json_filename': output_json_filename,
                'analysis_id': analysis_id,
                'original_filename': file.filename
            })
    
    return jsonify({'tasks': tasks}), 202

def process_pdf_v2(
    file_path: str, 
    output_json_path: str, 
    unique_filename: str, 
    output_json_filename: str,
    analysis_id: str,
    original_filename: str
):
    """Обработка PDF"""
    status_data_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{unique_filename}_status.json")
    
    def update_status(message: str, progress: int):
        status_entry = {'status': 'processing', 'message': message, 'progress': progress}
        try:
            with open(status_data_path, 'w', encoding='utf-8') as f:
                json.dump(status_entry, f)
        except IOError as e:
            print(f"Ошибка записи статуса в файл {status_data_path}: {e}")

    try:
        update_status("Инициализация ИИ-движка...", 5)
        time.sleep(0.5)
        
        print(f"📂 Конвертация PDF '{file_path}' в изображения (DPI={DPI})...")
        update_status("Конвертация PDF в изображения...", 10)
        pages = convert_from_path(file_path, dpi=DPI)
        print(f"🖼️ Преобразовано {len(pages)} страниц.")
        
        update_status("Анализ структуры документа...", 15)
        time.sleep(0.3)
        
        print("\n🔍 --- ЗАПУСК УЛУЧШЕННОГО OCR ДВИЖКА ---")
        structured_data = extract_structured_data_v2(
            pages, 
            original_filename,
            status_callback=update_status
        )
        
        print("\n📊 Извлеченные данные:")
        for key, value in structured_data.items():
            if key not in ["metrics", "confidence_scores", "summary", "document_type", "file_name", "validation_status", "extraction_accuracy"]: # Exclude non-core fields from general display
                print(f"  {key}: {value}")
        
        print(f"\n📝 Краткое описание: {structured_data.get('summary', '—')}") # Display summary
        print(f"\n✅ Точность извлечения: {structured_data['extraction_accuracy'] * 100:.1f}%")
        print("--- КОНЕЦ ИЗВЛЕЧЕНИЯ ДАННЫХ ---\n")
        
        update_status("Финализация результатов...", 98)
        
        # Сохраняем результат
        with open(output_json_path, 'w', encoding='utf-8') as f:
            json.dump(structured_data, f, ensure_ascii=False, indent=2, sort_keys=True)
        
        # Сохраняем в историю
        save_to_history(structured_data, analysis_id)
        
        # Для API статуса
        result_data_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{unique_filename}_result.json")
        with open(result_data_path, 'w', encoding='utf-8') as f:
            json.dump(structured_data, f, ensure_ascii=False, indent=2)
        
        update_status("Завершено!", 100)
        time.sleep(0.5)
        
    except Exception as e:
        print(f"❌ Ошибка при обработке файла {file_path}: {e}")
        traceback.print_exc()
        error_message = str(e)
        
        # Сохраняем сообщение об ошибке
        error_data_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{unique_filename}_error.json")
        try:
            with open(error_data_path, 'w', encoding='utf-8') as f:
                json.dump({'error': error_message}, f)
        except IOError as io_err:
            print(f"Ошибка записи ошибки в файл {error_data_path}: {io_err}")
        
        # Попытка обновить статус, если он еще не установлен или не завершен
        update_status(f"Ошибка: {error_message}", 0)


@app.route('/status/<task_id>')
def get_status(task_id):
    result_data_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{task_id}_result.json")
    error_data_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{task_id}_error.json")
    status_data_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{task_id}_status.json")
    
    if os.path.exists(result_data_path):
        try:
            with open(result_data_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'status': 'completed', 'data': data})
        except Exception as e:
            return jsonify({'status': 'error', 'data': {'error': f"Ошибка чтения файла результата: {e}"}})
            
    elif os.path.exists(error_data_path):
        try:
            with open(error_data_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'status': 'error', 'data': data})
        except Exception as e:
            return jsonify({'status': 'error', 'data': {'error': f"Ошибка чтения файла ошибки: {e}"}})
            
    elif os.path.exists(status_data_path):
        try:
            with open(status_data_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify(data)
        except Exception as e:
            return jsonify({'status': 'error', 'data': {'error': f"Ошибка чтения файла статуса: {e}"}})
            
    else:
        return jsonify({'status': 'unknown', 'message': 'Статус файла не найден.'})

@app.route('/status/<file_id>')
def check_status(file_id):
    """
    Возвращает текущий статус обработки файла
    """
    if file_id not in processing_queue:
        return jsonify({'error': 'Файл не найден'}), 404
    
    status_info = processing_queue[file_id]
    response = {
        'status': status_info['status'],
        'progress': status_info.get('progress', 0),
        'filename': status_info.get('filename', '')
    }
    
    if status_info['status'] == 'completed':
        response['result'] = status_info.get('result')
    elif status_info['status'] == 'error':
        response['error'] = status_info.get('error')
    
    return jsonify(response)

@app.route('/download/<filename>')
def download_file(filename):
    # This route is likely intended for the old method of returning status/results.
    # The new code introduces specific download routes for JSON and Excel.
    # Keeping it for backward compatibility or if it serves a different purpose.
    # This route is not directly called by the updated front-end logic.
    # It's best practice to use absolute paths or ensure the file path is constructed safely.
    file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if os.path.exists(file_path):
        return send_file(
            file_path, 
            as_attachment=True,
            download_name=f"contract_data_{filename}"
        )
    else:
        return jsonify({'error': 'Файл не найден'}), 404

@app.route('/history')
def get_history():
    """
    Возвращает историю всех анализов
    """
    history = load_history()
    return jsonify(history)

@app.route('/download/json/<analysis_id>')
def download_json(analysis_id):
    history_file = os.path.join(HISTORY_FOLDER, f"{analysis_id}.json")
    if os.path.exists(history_file):
        return send_file(
            history_file,
            as_attachment=True,
            download_name=f"contract_data_{analysis_id}.json",
            mimetype='application/json'
        )
    return jsonify({'error': 'Файл не найден'}), 404

@app.route('/download/excel/<analysis_id>')
def download_excel(analysis_id):
    history_file = os.path.join(HISTORY_FOLDER, f"{analysis_id}.json")
    if os.path.exists(history_file):
        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Передача данных в функцию create_excel_file, которая теперь вызывает export_to_excel
            excel_file = create_excel_file(data)
            
            return send_file(
                excel_file,
                as_attachment=True,
                download_name=f"contract_data_{analysis_id}.xlsx",
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        except FileNotFoundError:
            return jsonify({'error': 'Файл истории не найден'}), 404
        except json.JSONDecodeError:
            return jsonify({'error': 'Ошибка декодирования JSON файла истории'}), 500
        except Exception as e:
            print(f"Error generating Excel: {e}")
            return jsonify({'error': 'Ошибка при генерации Excel файла'}), 500
            
    return jsonify({'error': 'Файл не найден'}), 404

@app.route('/download/docx/<analysis_id>')
def download_docx(analysis_id):
    """Скачивание результатов анализа в формате DOCX"""
    history_file = os.path.join(HISTORY_FOLDER, f"{analysis_id}.json")
    if os.path.exists(history_file):
        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Генерация DOCX файла
            docx_file = export_to_docx(data)
            
            return send_file(
                docx_file,
                as_attachment=True,
                download_name=f"contract_data_{analysis_id}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except FileNotFoundError:
            return jsonify({'error': 'Файл истории не найден'}), 404
        except json.JSONDecodeError:
            return jsonify({'error': 'Ошибка декодирования JSON файла истории'}), 500
        except Exception as e:
            print(f"Error generating DOCX: {e}")
            return jsonify({'error': 'Ошибка при генерации DOCX файла'}), 500
            
    return jsonify({'error': 'Файл не найден'}), 404

@app.route('/history/<analysis_id>')
def get_analysis(analysis_id):
    history_file = os.path.join(HISTORY_FOLDER, f"{analysis_id}.json")
    if os.path.exists(history_file):
        with open(history_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return jsonify(data)
    return jsonify({'error': 'Анализ не найден'}), 404


# --- HTML Template ---
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>OCR 2.0 — Банковская система анализа</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

        :root {
            --primary: #1e3a8a;
            --primary-light: #3b82f6;
            --accent: #10b981;
            --accent-hover: #059669;
            --bg-main: #f8fafc;
            --bg-card: #ffffff;
            --text-primary: #0f172a;
            --text-secondary: #64748b;
            --border: #e2e8f0;
            --shadow-md: 0 4px 12px rgba(0,0,0,0.1);
            --radius: 12px;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Inter', sans-serif;
            background: var(--bg-main);
            color: var(--text-primary);
            line-height: 1.6;
        }

        .app-container {
            display: flex;
            min-height: 100vh;
        }

        .sidebar {
            width: 320px;
            background: var(--bg-card);
            border-right: 1px solid var(--border);
            position: fixed;
            height: 100vh;
            overflow-y: auto;
        }

        .sidebar-header {
            padding: 24px;
            background: linear-gradient(135deg, var(--primary) 0%, #1e40af 100%);
            color: white;
        }

        .sidebar-header h1 {
            font-size: 1.25rem;
            font-weight: 700;
            display: flex;
            align-items: center;
            gap: 10px;
        }

        .history-section {
            padding: 20px;
        }

        .history-section h2 {
            font-size: 0.875rem;
            font-weight: 600;
            text-transform: uppercase;
            color: var(--text-secondary);
            margin-bottom: 16px;
        }

        .history-list {
            list-style: none;
        }

        .history-item {
            padding: 14px 16px;
            margin-bottom: 8px;
            background: var(--bg-main);
            border: 1px solid var(--border);
            border-radius: 8px;
            cursor: pointer;
            transition: all 0.3s;
        }

        .history-item:hover {
            background: #f1f5f9;
            border-color: var(--primary-light);
            transform: translateX(4px);
        }

        .history-item-name {
            font-size: 0.875rem;
            font-weight: 500;
            margin-bottom: 6px;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }

        .history-item-meta {
            display: flex;
            justify-content: space-between;
            font-size: 0.75rem;
            color: var(--text-secondary);
        }

        .status-badge {
            padding: 2px 8px;
            border-radius: 12px;
            font-size: 0.7rem;
            font-weight: 600;
        }

        .status-success {
            background: #d1fae5;
            color: #065f46;
        }

        .status-partial {
            background: #fef3c7;
            color: #92400e;
        }

        .main-content {
            flex: 1;
            margin-left: 320px;
            padding: 32px;
            max-width: 1400px;
        }

        .header {
            text-align: center;
            margin-bottom: 40px;
        }

        .header h1 {
            font-size: 2.5rem;
            font-weight: 800;
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 12px;
        }

        .upload-card {
            background: var(--bg-card);
            border-radius: var(--radius);
            padding: 48px;
            box-shadow: var(--shadow-md);
            margin-bottom: 32px;
        }

        .drop-zone {
            border: 2px dashed var(--border);
            border-radius: var(--radius);
            padding: 64px 32px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s;
            background: var(--bg-main);
        }

        .drop-zone:hover {
            border-color: var(--primary);
            background: #eff6ff;
        }

        .drop-zone-icon {
            font-size: 3rem;
            color: var(--primary);
            margin-bottom: 16px;
        }

        .btn {
            padding: 12px 28px;
            border: none;
            border-radius: 8px;
            font-size: 1rem;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s;
            display: inline-flex;
            align-items: center;
            gap: 8px;
        }

        .btn-primary {
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-light) 100%);
            color: white;
            box-shadow: 0 4px 12px rgba(30, 58, 138, 0.3);
        }

        .btn-primary:hover {
            transform: translateY(-2px);
        }

        .btn-success {
            background: var(--accent);
            color: white;
            margin: 0 8px;
        }

        .btn-success:hover {
            background: var(--accent-hover);
        }

        .file-list {
            margin-top: 20px;
        }

        .file-item {
            display: flex;
            align-items: center;
            justify-content: space-between;
            padding: 12px 16px;
            background: #eff6ff;
            border: 1px solid #bfdbfe;
            border-radius: 8px;
            margin-bottom: 8px;
        }

        .results-container {
            display: flex;
            flex-direction: column;
            gap: 24px;
        }

        .result-document {
            background: var(--bg-card);
            border-radius: var(--radius);
            padding: 24px;
            box-shadow: var(--shadow-md);
            border: 1px solid var(--border);
        }

        .result-document-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
            padding-bottom: 16px;
            border-bottom: 2px solid var(--bg-main);
        }

        .result-document-title {
            font-size: 1.25rem;
            font-weight: 700;
            color: var(--text-primary);
        }

        .result-document-actions {
            display: flex;
            gap: 8px;
        }

        .results-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
        }

        .result-card {
            background: var(--bg-main);
            border-radius: 8px;
            padding: 20px;
            border: 1px solid var(--border);
        }

        .result-card-header {
            display: flex;
            align-items: center;
            gap: 10px;
            margin-bottom: 16px;
            padding-bottom: 12px;
            border-bottom: 2px solid var(--border);
        }

        .result-card-icon {
            font-size: 1.25rem;
            color: var(--primary);
        }

        .result-card-title {
            font-size: 1rem;
            font-weight: 700;
        }

        .result-list {
            list-style: none;
        }

        .result-item {
            display: flex;
            justify-content: space-between;
            padding: 10px 0;
            border-bottom: 1px solid var(--border);
        }

        .result-item:last-child {
            border-bottom: none;
        }

        .result-label {
            font-size: 0.875rem;
            font-weight: 600;
            color: var(--text-secondary);
        }

        .result-value {
            font-size: 0.875rem;
            color: var(--text-primary);
            text-align: right;
        }

        .confidence-bar {
            width: 100%;
            height: 6px;
            background: #e2e8f0;
            border-radius: 3px;
            overflow: hidden;
            margin-top: 6px;
        }

        .confidence-fill {
            height: 100%;
            background: linear-gradient(90deg, var(--accent) 0%, #34d399 100%);
            border-radius: 3px;
        }

        .spinner {
            width: 48px;
            height: 48px;
            border: 4px solid #e2e8f0;
            border-top: 4px solid var(--primary);
            border-radius: 50%;
            margin: 20px auto;
            animation: spin 1s linear infinite;
        }

        @keyframes spin {
            to { transform: rotate(360deg); }
        }

        .progress-message {
            text-align: center;
            font-size: 0.875rem;
            color: var(--text-secondary);
            margin-top: 12px;
        }

        .progress-bar {
            width: 100%;
            height: 6px;
            background: #e2e8f0;
            border-radius: 3px;
            overflow: hidden;
            margin-top: 8px;
        }

        .progress-fill {
            height: 100%;
            background: linear-gradient(90deg, var(--primary) 0%, var(--primary-light) 100%);
            transition: width 0.5s ease;
        }
    </style>
</head>
<body>
    <div class="app-container">
        <aside class="sidebar">
            <div class="sidebar-header">
                <h1><i class="fas fa-shield-alt"></i> OCR 2.0</h1>
                <p>Банковская система анализа</p>
            </div>
            <div class="history-section">
                <h2><i class="fas fa-history"></i> История анализа</h2>
                <ul class="history-list" id="historyList"></ul>
            </div>
        </aside>

        <main class="main-content">
            <header class="header">
                <h1>Система анализа банковских документов</h1>
                <p>Загрузите PDF-договоры для автоматического извлечения данных с использованием ИИ</p>
            </header>

            <section class="upload-card" id="uploadSection">
                <div class="drop-zone" id="dropZone">
                    <div class="drop-zone-icon">
                        <i class="fas fa-cloud-upload-alt"></i>
                    </div>
                    <div style="font-size: 1.125rem; margin-bottom: 8px;">Перетащите PDF-файлы сюда</div>
                    <div style="font-size: 0.875rem; color: var(--text-secondary); margin-bottom: 24px;">или нажмите кнопку для выбора нескольких файлов</div>
                    <button class="btn btn-primary" id="browseBtn">
                        <i class="fas fa-folder-open"></i>
                        Выбрать файлы
                    </button>
                    <input type="file" id="fileInput" accept=".pdf" multiple hidden>
                </div>
                <div class="file-list" id="fileList" style="display: none;"></div>
                <div style="text-align: center; margin-top: 32px;">
                    <button class="btn btn-primary" id="uploadBtn" disabled style="font-size: 1.125rem; padding: 14px 32px;">
                        <i class="fas fa-rocket"></i>
                        Начать анализ
                    </button>
                </div>
            </section>

            <div class="results-container" id="resultsContainer"></div>
        </main>
    </div>

    <script>
        let selectedFiles = [];
        let processingTasks = new Map();

        const dropZone = document.getElementById('dropZone');
        const fileInput = document.getElementById('fileInput');
        const browseBtn = document.getElementById('browseBtn');
        const fileList = document.getElementById('fileList');
        const uploadBtn = document.getElementById('uploadBtn');
        const resultsContainer = document.getElementById('resultsContainer');
        const historyList = document.getElementById('historyList');

        loadHistory();

        browseBtn.addEventListener('click', () => fileInput.click());

        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length) {
                handleFilesSelect(Array.from(e.target.files));
            }
        });

        ['dragenter', 'dragover', 'dragleave', 'drop'].forEach(eventName => {
            dropZone.addEventListener(eventName, (e) => {
                e.preventDefault();
                e.stopPropagation();
            });
        });

        dropZone.addEventListener('drop', (e) => {
            const files = Array.from(e.dataTransfer.files).filter(f => f.type === 'application/pdf');
            if (files.length) {
                handleFilesSelect(files);
            }
        });

        function handleFilesSelect(files) {
            selectedFiles = files;
            displayFileList();
            uploadBtn.disabled = false;
        }

        function displayFileList() {
            fileList.innerHTML = '';
            fileList.style.display = 'block';
            
            selectedFiles.forEach((file, index) => {
                const div = document.createElement('div');
                div.className = 'file-item';
                div.innerHTML = `
                    <div style="display: flex; align-items: center; gap: 12px;">
                        <i class="fas fa-file-pdf" style="font-size: 1.5rem; color: var(--primary);"></i>
                        <span style="font-weight: 500;">${file.name}</span>
                    </div>
                    <button class="btn" onclick="removeFile(${index})" style="background: var(--text-secondary); color: white; padding: 6px 12px;">
                        <i class="fas fa-times"></i>
                    </button>
                `;
                fileList.appendChild(div);
            });
        }

        window.removeFile = function(index) {
            selectedFiles.splice(index, 1);
            if (selectedFiles.length === 0) {
                fileList.style.display = 'none';
                uploadBtn.disabled = true;
            } else {
                displayFileList();
            }
        };

        uploadBtn.addEventListener('click', async () => {
            if (selectedFiles.length === 0) return;

            const formData = new FormData();
            selectedFiles.forEach(file => {
                formData.append('files', file);
            });

            uploadBtn.disabled = true;

            try {
                const response = await fetch('/upload', {
                    method: 'POST',
                    body: formData
                });

                if (response.ok) {
                    const data = await response.json();
                    data.tasks.forEach(task => {
                        createResultCard(task);
                        pollStatus(task);
                    });
                    
                    selectedFiles = [];
                    fileList.style.display = 'none';
                    fileInput.value = '';
                } else {
                    const errorData = await response.json();
                    alert('Ошибка загрузки файлов: ' + (errorData.error || response.statusText));
                    uploadBtn.disabled = false;
                }
            } catch (error) {
                console.error('Ошибка:', error);
                alert('Ошибка сети');
                uploadBtn.disabled = false;
            }
        });

        function createResultCard(task) {
            const card = document.createElement('div');
            card.className = 'result-document';
            card.id = `result-${task.task_id}`;
            card.innerHTML = `
                <div class="result-document-header">
                    <div class="result-document-title">
                        <i class="fas fa-file-pdf" style="color: var(--primary); margin-right: 8px;"></i>
                        ${task.original_filename}
                    </div>
                    <div class="result-document-actions" id="actions-${task.task_id}" style="display: none;">
                        <button class="btn btn-success" onclick="downloadJSON('${task.analysis_id}')">
                            <i class="fas fa-download"></i>
                            JSON
                        </button>
                        <button class="btn btn-success" onclick="downloadExcel('${task.analysis_id}')">
                            <i class="fas fa-file-excel"></i>
                            Excel
                        </button>
                        <button class="btn btn-success" onclick="downloadDOCX('${task.analysis_id}')">
                            <i class="fas fa-file-word"></i>
                            DOCX
                        </button>
                    </div>
                </div>
                <div id="content-${task.task_id}">
                    <div class="spinner"></div>
                    <div class="progress-message" id="progress-msg-${task.task_id}">Инициализация...</div>
                    <div class="progress-bar">
                        <div class="progress-fill" id="progress-${task.task_id}" style="width: 0%"></div>
                    </div>
                </div>
            `;
            resultsContainer.appendChild(card);
        }

        async function pollStatus(task) {
            const checkStatus = async () => {
                try {
                    const response = await fetch(`/status/${task.task_id}`);
                    const data = await response.json();

                    if (data.status === 'completed') {
                        displayResults(task, data.data);
                        loadHistory(); // Reload history after completion
                        uploadBtn.disabled = false; // Re-enable upload button
                    } else if (data.status === 'error') {
                        displayError(task, data.data.error || 'Неизвестная ошибка');
                        loadHistory(); // Reload history even on error
                        uploadBtn.disabled = false; // Re-enable upload button
                    } else if (data.status === 'processing') {
                        const progress = data.progress || 0;
                        document.getElementById(`progress-${task.task_id}`).style.width = `${progress}%`;
                        document.getElementById(`progress-msg-${task.task_id}`).textContent = data.message || 'Обработка...';
                        setTimeout(checkStatus, 800); // Check again after 800ms
                    } else { // Status is 'queued' or 'unknown'
                        // Keep checking, maybe with a slightly longer interval for initial states
                        setTimeout(checkStatus, 1500); 
                    }
                } catch (error) {
                    console.error('Ошибка проверки статуса:', error);
                    document.getElementById(`progress-msg-${task.task_id}`).textContent = 'Ошибка связи с сервером';
                    document.getElementById(`progress-${task.task_id}`).style.backgroundColor = 'red'; // Indicate error visually
                    // Stop polling on persistent network errors
                    setTimeout(checkStatus, 5000); // Try again after a longer interval
                }
            };

            checkStatus();
        }

        function displayResults(task, data) {
            const content = document.getElementById(`content-${task.task_id}`);
            const actions = document.getElementById(`actions-${task.task_id}`);
            
            actions.style.display = 'flex';
            
            content.innerHTML = `
                <div class="results-grid">
                    <div class="result-card">
                        <div class="result-card-header">
                            <i class="fas fa-info-circle result-card-icon"></i>
                            <h3 class="result-card-title">Общая информация</h3>
                        </div>
                        <ul class="result-list">
                            <li class="result-item">
                                <span class="result-label">Тип документа</span>
                                <span class="result-value">${data.document_type || 'contract'}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Имя файла</span>
                                <span class="result-value">${data.file_name}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Статус валидации</span>
                                <span class="result-value">${data.validation_status || 'valid'}</span>
                            </li>
                        </ul>
                    </div>

                    <div class="result-card" style="grid-column: 1 / -1;">
                        <div class="result-card-header">
                            <i class="fas fa-file-alt result-card-icon"></i>
                            <h3 class="result-card-title">Краткое описание договора</h3>
                        </div>
                        <div style="padding: 16px; background: white; border-radius: 8px; line-height: 1.8; color: var(--text-primary); font-size: 0.9rem;">
                            ${data.summary || 'Описание недоступно'}
                        </div>
                    </div>

                    <div class="result-card">
                        <div class="result-card-header">
                            <i class="fas fa-file-contract result-card-icon"></i>
                            <h3 class="result-card-title">Данные договора</h3>
                        </div>
                        <ul class="result-list">
                            <li class="result-item">
                                <span class="result-label">Номер договора</span>
                                <span class="result-value">${data.contract_number || '—'}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Дата подписания</span>
                                <span class="result-value">${data.sign_date || '—'}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Дата окончания</span>
                                <span class="result-value">${data.expiry_date || '—'}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Продавец</span>
                                <span class="result-value">${data.seller || '—'}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Покупатель</span>
                                <span class="result-value">${data.buyer || '—'}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Сумма</span>
                                <span class="result-value">${data.amount || '—'}</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Валюта</span>
                                <span class="result-value">${data.currency || '—'}</span>
                            </li>
                        </ul>
                    </div>
                    <div class="result-card">
                        <div class="result-card-header">
                            <i class="fas fa-chart-line result-card-icon"></i>
                            <h3 class="result-card-title">Метрики качества</h3>
                        </div>
                        <ul class="result-list">
                            <li class="result-item">
                                <span class="result-label">Точность извлечения</span>
                                <span class="result-value">${(data.metrics.field_level_accuracy * 100).toFixed(1)}%</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Полное совпадение</span>
                                <span class="result-value">${(data.metrics.exact_match * 100).toFixed(1)}%</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Валидность JSON</span>
                                <span class="result-value">${(data.metrics.json_validity * 100).toFixed(1)}%</span>
                            </li>
                            <li class="result-item">
                                <span class="result-label">Согласованность схемы</span>
                                <span class="result-value">${(data.metrics.schema_consistency * 100).toFixed(1)}%</span>
                            </li>
                        </ul>
                    </div>
                    <div class="result-card">
                        <div class="result-card-header">
                            <i class="fas fa-brain result-card-icon"></i>
                            <h3 class="result-card-title">Оценки уверенности ИИ</h3>
                        </div>
                        <ul class="result-list">
                            ${Object.entries(data.confidence_scores || {}).map(([key, value]) => {
                                const labels = {
                                    'contract_number': 'Номер договора',
                                    'sign_date': 'Дата подписания',
                                    'expiry_date': 'Дата окончания',
                                    'seller': 'Продавец',
                                    'buyer': 'Покупатель',
                                    'amount': 'Сумма',
                                    'currency': 'Валюта'
                                };
                                return `
                                    <li class="result-item">
                                        <div style="flex: 1;">
                                            <div style="display: flex; justify-content: space-between; margin-bottom: 6px;">
                                                <span class="result-label">${labels[key] || key}</span>
                                                <span class="result-value">${(value * 100).toFixed(0)}%</span>
                                            </div>
                                            <div class="confidence-bar">
                                                <div class="confidence-fill" style="width: ${value * 100}%"></div>
                                            </div>
                                        </div>
                                    </li>
                                `;
                            }).join('')}
                        </ul>
                    </div>
                </div>
            `;
        }

        function displayError(task, error) {
            const content = document.getElementById(`content-${task.task_id}`);
            content.innerHTML = `
                <div style="text-align: center; padding: 40px; color: #ef4444;">
                    <i class="fas fa-exclamation-triangle" style="font-size: 3rem; margin-bottom: 16px;"></i>
                    <h3 style="font-size: 1.5rem; margin-bottom: 12px;">Ошибка обработки</h3>
                    <p>${error}</p>
                </div>
            `;
            // Ensure actions buttons are visible even on error to allow downloads if partially processed
            const actions = document.getElementById(`actions-${task.task_id}`);
            if (actions) actions.style.display = 'flex';
        }

        window.downloadJSON = function(analysisId) {
            window.location.href = `/download/json/${analysisId}`;
        };

        window.downloadExcel = function(analysisId) {
            window.location.href = `/download/excel/${analysisId}`;
        };

        window.downloadDOCX = function(analysisId) {
            window.location.href = `/download/docx/${analysisId}`;
        };

        async function loadHistory() {
            try {
                const response = await fetch('/history');
                if (!response.ok) {
                    throw new Error(`HTTP error! status: ${response.status}`);
                }
                const history = await response.json();
                
                const historyListElement = document.getElementById('historyList');
                historyListElement.innerHTML = ''; // Clear existing list
                
                if (history.length === 0) {
                    historyListElement.innerHTML = '<li style="text-align: center; padding: 40px; color: var(--text-secondary);">Анализов пока нет</li>';
                    return;
                }
                
                history.forEach(item => {
                    const li = document.createElement('li');
                    li.className = 'history-item';
                    li.dataset.analysisId = item.analysis_id; // Store ID for potential future use
                    
                    const date = new Date(item.timestamp);
                    // Format date to be user-friendly
                    const dateStr = date.toLocaleString('ru-RU', {
                        day: '2-digit',
                        month: '2-digit',
                        year: 'numeric',
                        hour: '2-digit',
                        minute: '2-digit'
                    });
                    
                    const statusClass = item.validation_status === 'valid' ? 'success' : 'partial';
                    const statusText = item.validation_status === 'valid' ? 'Успешно' : 'Частично';
                    
                    li.innerHTML = `
                        <div class="history-item-name">${item.file_name}</div>
                        <div class="history-item-meta">
                            <span><i class="fas fa-clock"></i> ${dateStr}</span>
                            <span class="status-badge status-${statusClass}">${statusText}</span>
                        </div>
                    `;
                    
                    li.addEventListener('click', () => viewHistoryItem(item.analysis_id));
                    
                    historyListElement.appendChild(li);
                });
            } catch (error) {
                console.error('Ошибка загрузки истории:', error);
                const historyListElement = document.getElementById('historyList');
                historyListElement.innerHTML = '<li style="text-align: center; padding: 40px; color: red;">Ошибка загрузки истории</li>';
            }
        }
        
        async function viewHistoryItem(analysisId) {
            try {
                const response = await fetch(`/history/${analysisId}`);
                if (!response.ok) {
                    throw new Error(`HTTP error! status: ${response.status}`);
                }
                const data = await response.json();
                
                // Clear results container and show this analysis
                resultsContainer.innerHTML = '';
                
                // Create a task object to reuse existing display function
                const task = {
                    task_id: analysisId,
                    analysis_id: analysisId,
                    original_filename: data.file_name
                };
                
                createResultCard(task);
                displayResults(task, data);
                
                // Scroll to results
                resultsContainer.scrollIntoView({ behavior: 'smooth' });
            } catch (error) {
                console.error('Ошибка загрузки анализа:', error);
                alert('Ошибка загрузки данных анализа');
            }
        }
    </script>
</body>
</html>
'''

if __name__ == '__main__':
    print("=" * 80)
    print("🚀 OCR 2.0 — БАНКОВСКАЯ СИСТЕМА АНАЛИЗА ДОКУМЕНТОВ")
    print("=" * 80)
    print("📊 Возможности:")
    print("   • Множественная загрузка PDF документов")
    print("   • Последовательная обработка с отображением прогресса")
    print("   • Извлечение ключевых полей договора (номер, даты, стороны, сумма, валюта)")
    print("   • Генерация краткого описания документа с помощью ИИ")
    print("   • Экспорт результатов в JSON, Excel и DOCX")
    print("   • Индивидуальные кнопки скачивания для каждого документа")
    print("   • Отображение истории анализов")
    print("=" * 80)
    print("🌐 Веб-интерфейс: http://localhost:5000")
    print("=" * 80)
    app.run(debug=True, host='0.0.0.0', port=5000)